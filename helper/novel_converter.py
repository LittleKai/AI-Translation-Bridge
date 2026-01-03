import os
import re
import csv
import glob
from pathlib import Path
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import zipfile
import xml.etree.ElementTree as ET
from urllib.parse import unquote
import html
import re
import pandas as pd

def natural_sort_key(s):
    """
    Sort file names naturally by numbers
    """
    return [int(text) if text.isdigit() else text.lower()
            for text in re.split(r'(\d+)', s)]


def detect_encoding(file_path):
    """
    Detect file encoding
    """
    encodings = ['utf-8', 'utf-16', 'utf-16-le', 'utf-16-be', 'gb18030', 'gbk', 'big5', 'shift-jis', 'euc-jp', 'euc-kr']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
                return encoding
        except UnicodeDecodeError:
            continue
    return 'utf-8'


def read_txt_content(file_path):
    """
    Read TXT file content
    """
    try:
        encoding = detect_encoding(file_path)
        with open(file_path, 'r', encoding=encoding, errors='ignore') as file:
            content = file.read()
            content = content.replace('\r\n', '\n').replace('\r', '\n')
            content = content.replace('\u2028', '\n').replace('\u2029', '\n')
            return content.split('\n')
    except Exception as e:
        return []


def read_docx_content(file_path):
    """
    Read DOCX file content
    """
    try:
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in paragraphs:
                        paragraphs.append(text)
        return paragraphs
    except PackageNotFoundError:
        return []
    except Exception as e:
        return []


def parse_container_xml(epub_zip):
    """
    Parse container.xml to find main OPF file
    """
    try:
        container_content = epub_zip.read('META-INF/container.xml')
        container_root = ET.fromstring(container_content)
        namespaces = {'container': 'urn:oasis:names:tc:opendocument:xmlns:container'}
        rootfile = container_root.find('.//container:rootfile', namespaces)
        if rootfile is not None:
            return rootfile.get('full-path')
        rootfile = container_root.find('.//rootfile')
        if rootfile is not None:
            return rootfile.get('full-path')
    except Exception:
        pass
    return None


def parse_opf_spine(epub_zip, opf_path):
    """
    Parse OPF file to get spine order
    """
    try:
        opf_content = epub_zip.read(opf_path)
        opf_root = ET.fromstring(opf_content)
        namespaces = {
            'opf': 'http://www.idpf.org/2007/opf',
            'dc': 'http://purl.org/dc/elements/1.1/'
        }
        manifest = {}
        manifest_items = opf_root.findall('.//opf:item', namespaces)
        if not manifest_items:
            manifest_items = opf_root.findall('.//item')
        for item in manifest_items:
            item_id = item.get('id')
            href = item.get('href')
            if item_id and href:
                href = unquote(href)
                if not href.startswith('/'):
                    opf_dir = os.path.dirname(opf_path)
                    if opf_dir:
                        href = os.path.join(opf_dir, href).replace('\\', '/')
                manifest[item_id] = href
        ordered_files = []
        spine_items = opf_root.findall('.//opf:itemref', namespaces)
        if not spine_items:
            spine_items = opf_root.findall('.//itemref')
        for itemref in spine_items:
            idref = itemref.get('idref')
            if idref and idref in manifest:
                file_path = manifest[idref]
                if file_path.lower().endswith(('.html', '.xhtml', '.htm')):
                    ordered_files.append(file_path)
        return ordered_files
    except Exception:
        return []


def clean_html_text(text):
    """
    Clean HTML text and remove tags
    """
    text = re.sub(r'<br\s*/?>', '\n', text)
    text = re.sub(r'<p[^>]*>', '\n', text)
    text = re.sub(r'</p>', '', text)
    text = re.sub(r'<[^>]+>', '', text)
    text = html.unescape(text)
    text = text.strip()
    return text

def generate_output_path(input_path, language):
    """
    Generate output path based on input path
    """
    input_name = os.path.basename(input_path)
    if os.path.isfile(input_path):
        input_name = os.path.splitext(input_name)[0]

    if f"_{language}" in input_name:
        output_filename = f"{input_name}.csv"
    else:
        output_filename = f"{input_name}_{language}.csv"

    data_dir_pattern = r'(.*?[/\\]Data[/\\])'
    match = re.search(data_dir_pattern, input_path)

    if match:
        data_dir = match.group(1)
        output_dir = os.path.join(data_dir, "Output", language)
    else:
        output_dir = os.path.join(os.path.dirname(input_path), "Output", language)

    os.makedirs(output_dir, exist_ok=True)
    return os.path.join(output_dir, output_filename)

# helper/novel_converter.py
# Thêm hàm mới để xử lý output format

def save_to_file(all_rows, output_path, log_callback=None):
    """
    Save rows to CSV or Excel based on file extension
    """
    if not all_rows:
        if log_callback:
            log_callback("No content to write")
        return False

    try:
        # Create DataFrame
        df = pd.DataFrame(all_rows, columns=['id', 'text'])

        # Check output format by extension
        _, ext = os.path.splitext(output_path)
        ext = ext.lower()

        if ext == '.xlsx' or ext == '.xls':
            # Save as Excel
            df.to_excel(output_path, index=False, engine='openpyxl')
            if log_callback:
                log_callback(f"Saved as Excel file: {output_path}")
        else:
            # Save as CSV (default)
            with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
                csv_writer = csv.writer(csvfile)
                csv_writer.writerow(['id', 'text'])
                csv_writer.writerows(all_rows)
            if log_callback:
                log_callback(f"Saved as CSV file: {output_path}")

        if log_callback:
            log_callback(f"Processing completed! Total lines: {len(all_rows)}")

        return True
    except Exception as e:
        if log_callback:
            log_callback(f"Error writing file: {str(e)}")
        return False

def convert_to_csv(input_path, language, output_path=None, ruby_handling=None, log_callback=None):
    """
    Convert file or folder to CSV/Excel with id and text columns
    """
    if not os.path.exists(input_path):
        if log_callback:
            log_callback(f"Path does not exist: {input_path}")
        return False, None

    if output_path is None:
        from pathlib import Path
        output_dir = Path.home() / "Documents" / "AIBridge"
        output_dir.mkdir(parents=True, exist_ok=True)

        input_name = os.path.basename(input_path)
        if os.path.isfile(input_path):
            input_name = os.path.splitext(input_name)[0]

        if f"_{language}" in input_name:
            output_filename = f"{input_name}.csv"
        else:
            output_filename = f"{input_name}_{language}.csv"

        output_path = str(output_dir / output_filename)

    if log_callback:
        log_callback(f"Output will be saved to: {output_path}")

    if os.path.isdir(input_path):
        all_rows = process_folder_to_csv(input_path, language, ruby_handling, log_callback)
    else:
        all_rows, _ = process_file_to_csv(input_path, language, 1, ruby_handling, log_callback)

    # Use new save function that supports both CSV and Excel
    success = save_to_file(all_rows, output_path, log_callback)

    if success:
        return True, output_path
    else:
        return False, None

def process_file_to_csv(file_path, language, current_id=1, ruby_handling=None, log_callback=None):
    """
    Process a single file and return list of rows
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    file_name = os.path.basename(file_path)

    if log_callback:
        log_callback(f"Processing file: {file_name}")

    if file_ext == '.txt':
        lines = read_txt_content(file_path)
    elif file_ext == '.docx':
        lines = read_docx_content(file_path)
    elif file_ext == '.epub':
        lines = read_epub_content(file_path, ruby_handling)
    else:
        if log_callback:
            log_callback(f"Unsupported file format: {file_ext}")
        return [], current_id

    rows = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        rows.append([current_id, line])
        current_id += 1

    if log_callback:
        log_callback(f"Processed {len(rows)} lines from file")

    return rows, current_id

def process_folder_to_csv(folder_path, language, ruby_handling=None, log_callback=None):
    """
    Process all supported files in folder
    """
    if log_callback:
        log_callback(f"Searching for files in folder: {folder_path}")

    txt_files = glob.glob(os.path.join(folder_path, "*.txt"))
    docx_files = glob.glob(os.path.join(folder_path, "*.docx"))
    epub_files = glob.glob(os.path.join(folder_path, "*.epub"))

    all_files = txt_files + docx_files + epub_files

    if not all_files:
        if log_callback:
            log_callback(f"No supported files found in folder")
        return []

    all_files.sort(key=lambda x: natural_sort_key(os.path.basename(x)))

    if log_callback:
        log_callback(f"Found {len(all_files)} file(s): TXT={len(txt_files)}, DOCX={len(docx_files)}, EPUB={len(epub_files)}")

    all_rows = []
    current_id = 1

    for file_path in all_files:
        rows, current_id = process_file_to_csv(file_path, language, current_id, ruby_handling, log_callback)
        all_rows.extend(rows)

    return all_rows

def read_epub_content(file_path, ruby_handling=None):
    """
    Read EPUB file content with advanced HTML processing and image position tracking
    """
    try:
        paragraphs = []

        with zipfile.ZipFile(file_path, 'r') as epub_zip:
            file_list = epub_zip.namelist()

            opf_path = parse_container_xml(epub_zip)

            if opf_path:
                ordered_files = parse_opf_spine(epub_zip, opf_path)
                html_files = []

                for file_path_in_epub in ordered_files:
                    candidates = [
                        file_path_in_epub,
                        file_path_in_epub.lstrip('/'),
                        file_path_in_epub.replace('/', os.sep),
                        os.path.basename(file_path_in_epub)
                    ]

                    for candidate in candidates:
                        if candidate in file_list:
                            html_files.append(candidate)
                            break
                    else:
                        basename = os.path.basename(file_path_in_epub)
                        for f in file_list:
                            if os.path.basename(f) == basename and f.endswith(('.html', '.xhtml', '.htm')):
                                html_files.append(f)
                                break
            else:
                html_files = [f for f in file_list
                              if f.endswith(('.html', '.xhtml', '.htm'))
                              and not f.startswith('__MACOSX')]
                html_files.sort(key=natural_sort_key)

            for html_file in html_files:
                try:
                    content = epub_zip.read(html_file)

                    if isinstance(content, bytes):
                        content = content.decode('utf-8', errors='replace')

                    extracted_lines = extract_content_from_html(content, ruby_handling)
                    paragraphs.extend(extracted_lines)

                except Exception:
                    continue

        return paragraphs

    except Exception:
        return []


def is_hiragana(text):
    """
    Check if text is hiragana
    """
    return bool(re.match(r'^[\u3040-\u309F]+$', text))

def is_katakana(text):
    """
    Check if text is katakana
    """
    return bool(re.match(r'^[\u30A0-\u30FF]+$', text))

def is_css_content(text):
    """
    Check if text is CSS/style content
    """
    css_patterns = [
        r'^\s*[a-zA-Z#\.\[\]]+\s*\{\s*[^}]*\}\s*$',
        r'^\s*[a-zA-Z\-]+\s*:\s*[^;]+;\s*$',
        r'^\s*@[a-zA-Z\-]+\s*[^{]*\{\s*[^}]*\}\s*$',
        r'^\s*[a-zA-Z\-]+\s*:\s*[^;]+\s*$',
        r'.*\{\s*(padding|margin|text-align|display|list-style)[^}]*\}',
    ]

    for pattern in css_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            return True

    css_keywords = ['padding', 'margin', 'text-align', 'display', 'font-size',
                    'color', 'background', 'border', 'width', 'height', 'position',
                    'float', 'clear', 'list-style-type', 'nav#landmarks']

    for keyword in css_keywords:
        if keyword in text.lower() and ':' in text:
            return True

    if 'http://www.w3.org/' in text and 'DTD' in text:
        return True

    return False

def is_html_tag_or_attribute(text):
    """
    Check if text is HTML tag, attribute, or XML declaration
    Should be filtered out from content
    """
    text = text.strip()

    if not text:
        return True

    html_patterns = [
        r'^xmlns\s*=',
        r'^\s*xmlns:',
        r'^class\s*=',
        r'^id\s*=',
        r'^style\s*=',
        r'^version\s*=',
        r'^width\s*=|^height\s*=',
        r'^viewBox\s*=',
        r'http://www\.w3\.org/',
        r'^\s*"[^"]*"\s*$',
    ]

    for pattern in html_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            return True

    standalone_tags = ['xmlns', 'class', 'style', 'script', 'head', 'body', 'meta', 'link']
    if text.lower() in standalone_tags:
        return True

    if '=' in text and ('"' in text or "'" in text):
        if not re.search(r'[\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FFF\uAC00-\uD7AF]', text):
            return True

    return False

def decode_html_entities(text):
    """
    Decode HTML entities to Unicode characters
    """
    decoded = html.unescape(text)

    pattern = r'&#x([0-9A-Fa-f]{4,6});'
    def replace_entity(match):
        try:
            return chr(int(match.group(1), 16))
        except:
            return match.group(0)

    decoded = re.sub(pattern, replace_entity, decoded)
    return decoded

def process_ruby_tags(content, ruby_handling):
    """
    Process ruby tags based on handling mode with complex ruby support
    """
    if ruby_handling == 'remove_all':
        content = re.sub(r'<ruby>(.*?)<rt>.*?</rt></ruby>', r'\1', content, flags=re.DOTALL)
        content = re.sub(r'<ruby>(.*?)</ruby>', r'\1', content, flags=re.DOTALL)

    elif ruby_handling == 'remove_hiragana':
        def replace_complex_ruby(match):
            ruby_content = match.group(1)

            rb_rt_pairs = re.findall(r'<rb>(.*?)</rb>\s*<rt>(.*?)</rt>', ruby_content, re.DOTALL)

            if rb_rt_pairs:
                combined_rb = ""
                combined_rt = ""

                for rb_text, rt_text in rb_rt_pairs:
                    combined_rb += rb_text.strip()
                    combined_rt += rt_text.strip()

                if combined_rt and is_hiragana(combined_rt):
                    return combined_rb
                else:
                    return f'{combined_rb}({combined_rt})' if combined_rt else combined_rb
            else:
                simple_match = re.search(r'(.*?)<rt>(.*?)</rt>', ruby_content, re.DOTALL)
                if simple_match:
                    main_text = simple_match.group(1).strip()
                    rt_text = simple_match.group(2).strip()

                    if rt_text and is_hiragana(rt_text):
                        return main_text
                    else:
                        return f'{main_text}({rt_text})' if rt_text else main_text
                else:
                    return ruby_content

        content = re.sub(r'<ruby>(.*?)</ruby>', replace_complex_ruby, content, flags=re.DOTALL)

        def replace_simple_ruby(match):
            main_text = match.group(1).strip()
            rt_text = match.group(2).strip()

            if rt_text and is_hiragana(rt_text):
                return main_text
            else:
                return f'{main_text}({rt_text})' if rt_text else main_text

        content = re.sub(r'<ruby>(.*?)<rt>(.*?)</rt></ruby>', replace_simple_ruby, content, flags=re.DOTALL)

    elif ruby_handling == 'keep_all':
        def replace_complex_ruby_keep(match):
            ruby_content = match.group(1)

            rb_rt_pairs = re.findall(r'<rb>(.*?)</rb>\s*<rt>(.*?)</rt>', ruby_content, re.DOTALL)

            if rb_rt_pairs:
                combined_rb = ""
                combined_rt = ""

                for rb_text, rt_text in rb_rt_pairs:
                    combined_rb += rb_text.strip()
                    combined_rt += rt_text.strip()

                return f'{combined_rb}({combined_rt})' if combined_rt else combined_rb
            else:
                base_match = re.search(r'<rb>(.*?)</rb>', ruby_content, re.DOTALL)
                rt_match = re.search(r'<rt>(.*?)</rt>', ruby_content, re.DOTALL)

                if base_match:
                    base_text = base_match.group(1).strip()
                    if rt_match:
                        rt_text = rt_match.group(1).strip()
                        return f"{base_text}({rt_text})" if rt_text else base_text
                    return base_text

                return ruby_content

        content = re.sub(r'<ruby>(.*?)</ruby>', replace_complex_ruby_keep, content, flags=re.DOTALL)

        def replace_simple_ruby_keep(match):
            base_text = match.group(1).strip()
            rt_text = match.group(2).strip()
            return f"{base_text}({rt_text})" if rt_text else base_text

        content = re.sub(r'<ruby>([^<]+)<rt>([^<]*)</rt></ruby>', replace_simple_ruby_keep, content)

    return content

def merge_consecutive_ruby_readings(text):
    """
    Merge consecutive ruby readings into single parentheses
    Example: 情報収集妨害(ジャミング)煙幕(スモーク) → 情報収集妨害煙幕(ジャミングスモーク)
    """
    def merge_group(match):
        full_text = match.group(0)

        pairs = re.findall(r'([^\(]+)\(([^\)]+)\)', full_text)

        if len(pairs) >= 2:
            combined_text = ''.join([pair[0] for pair in pairs])
            combined_reading = ''.join([pair[1] for pair in pairs])
            return f'{combined_text}({combined_reading})'
        else:
            return full_text

    pattern = r'(?:[^\(]+\([^\)]+\)){2,}'
    text = re.sub(pattern, merge_group, text)

    return text

def extract_content_from_html(content, ruby_handling=None):
    """
    Extract content from HTML with advanced processing including image positions
    """
    if isinstance(content, bytes):
        try:
            content = content.decode('utf-8')
        except UnicodeDecodeError:
            content = content.decode('utf-8', errors='replace')

    content = re.sub(r'<!DOCTYPE[^>]*>', '', content, flags=re.IGNORECASE)
    content = re.sub(r'<\?xml[^>]*\?>', '', content, flags=re.IGNORECASE)
    content = re.sub(r'<!doctype[^>]*>', '', content, flags=re.IGNORECASE)

    content = re.sub(r'<!--.*?-->', '', content, flags=re.DOTALL)

    content = re.sub(r'<title[^>]*>.*?</title>', '', content, flags=re.DOTALL | re.IGNORECASE)

    if ruby_handling:
        content = process_ruby_tags(content, ruby_handling)

    result = []

    content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL | re.IGNORECASE)
    content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL | re.IGNORECASE)

    content = re.sub(r'<svg[^>]*>.*?</svg>', '', content, flags=re.DOTALL | re.IGNORECASE)
    content = re.sub(r'<svg[^>]*>', '', content, flags=re.IGNORECASE)

    content = re.sub(r'<img[^>]*>', '(img)', content, flags=re.IGNORECASE)

    content = re.sub(r'<br\s*/?>', '\n', content, flags=re.IGNORECASE)
    content = re.sub(r'</p>', '\n', content, flags=re.IGNORECASE)
    content = re.sub(r'</div>', '\n', content, flags=re.IGNORECASE)
    content = re.sub(r'</h[1-6]>', '\n', content, flags=re.IGNORECASE)
    content = re.sub(r'</li>', '\n', content, flags=re.IGNORECASE)
    content = re.sub(r'</td>', '\n', content, flags=re.IGNORECASE)
    content = re.sub(r'</th>', '\n', content, flags=re.IGNORECASE)
    content = re.sub(r'</blockquote>', '\n', content, flags=re.IGNORECASE)
    content = re.sub(r'</section>', '\n', content, flags=re.IGNORECASE)
    content = re.sub(r'</article>', '\n', content, flags=re.IGNORECASE)

    content = re.sub(r'<[^>]+>', '', content)

    content = decode_html_entities(content)

    lines = content.split('\n')
    for line in lines:
        cleaned = line.strip()

        if cleaned and not is_css_content(cleaned) and not is_html_tag_or_attribute(cleaned):
            if ruby_handling:
                cleaned = merge_consecutive_ruby_readings(cleaned)
            result.append(cleaned)

    return result